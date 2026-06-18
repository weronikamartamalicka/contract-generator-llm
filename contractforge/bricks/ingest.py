import docx
import logging
from typing import Protocol
from contractforge.contracts import RawDocument

class Parser(Protocol):
    def parse(self, path: str) -> RawDocument | None : ...

class DocxParser:
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def parse(self, path: str) -> RawDocument | None:
        self.logger.info("Starting to parse docx file %s", path)
        try:
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            self.logger.error("Error parsing docx file %s: %s", path, e)
        else:
            return RawDocument(
                source_format = "docx",
                text = text,
                source_name=path
            )